"""
Functions realted to scraping the data from an incoming cover sheet.
"""

import src.utility as utility
from src.constants import COVER_SHEET_DIRECTORY, HEADER_MAP

from docx import Document
from docx.text.paragraph import Paragraph
import pandas as pd
import os

def read_cover_sheet(document):
    """
    Returns a dictionary containing the data collected from the passed word document
    
    :param document: A Document object holding a performance cover sheet from a .docx file
    :return: A dictionary containing all of the relevant data scraped from the passed cover sheet document
    """
    data = {}
    header = None

    for block in utility.iter_block_items(document):
        if isinstance(block, Paragraph):    # handling of paragraphs
            if block.text in HEADER_MAP.keys():     # if the paragraph holds one of the headers used to indicate that text can be inputted from user
                header = block.text     # stores text as a header
        else:   # if block is Table object
            if header:  # if table comes directly after a recognized header
                table_data = get_list_from_table(block)

                # Storing of data under the mapped column name, resetting header to read more data
                data[HEADER_MAP[header]] = " ".join(table_data)
                header = None
            else:   # for tables that do not come directly after headers
                for row in block.rows:
                    row_text = []
                    for cell in row.cells:
                        for paragraph in cell.paragraphs: 
                            row_text.append(paragraph.text)     # collects all paragraphs in a given cell into a list

                    # storing of checkboxes in data field - continues until all checkboxes are retrieved from cell
                    while __checkbox_in_row(row_text):
                        # Retrieves index of the first checked and unchecked checkboxes
                        checked_index = __get_checkbox_index(row_text, True)
                        unchecked_index = __get_checkbox_index(row_text, False)

                        index = max(checked_index, unchecked_index)     # retrives the maximum index value, meaning index -1 is never selected
                        checkbox_value = 1 if index == checked_index else 0
                        row_text.pop(index)     # removes element with checkbox that is being added to data
                        column_title = row_text.pop(index)  # removes element in row directly after checkbox, which is assumed to be the title of the checkbox field

                        data[column_title] = checkbox_value   # stores a 1 if box is checked, else 0
    
    return data

# DOCUMENT RETRIEVAL METHODS

def process_cover_sheets(cover_sheets_list):
    """
    Creates a DataFrame object from a list of cover sheets
    
    :param cover_sheets_list: A list of Document objects, each of which was created from a separate .docx file containing a cover sheet
    :return: A DataFrame object with each row representing a cover sheet.
    """
    data = []
    for cover_sheet in cover_sheets_list:
        data.append(read_cover_sheet(cover_sheet))
        
    return pd.DataFrame(data)

def get_cover_sheets(path=None):
    """
    Returns a list of docx Document objects representing cover sheets, each of which are retrieved from the folder where cover sheets are stored.

    :param path: The path to the directory where cover sheet objects are stored. NOTE: This directory should only include cover sheets.
    :return: A list of docx Document objects representing cover sheets, each of which are retrieved from the folder where cover sheets are stored.
    """
    cover_sheets = []

    if path == None:
        path = COVER_SHEET_DIRECTORY    # retrieves cover sheets from the default directory specified by a constant if no path was passed

    try:
        for filename in os.listdir(path):  # retrieves all file names from cover sheet directory
            cover_sheets.append(Document(f"{path}{filename}"))     # creates Document objects for every cover sheet
    except FileNotFoundError as e:
        print(f"Unable to retrieve cover sheets from path {path}")
        return None

    return cover_sheets

# DATA READING METHODS

def get_list_from_table(table):
    """
    Given a passed table, returns a list with each element representing the data held in a cell of the table. Parses through the table from top to bottom, going left to right in each row.

    :param table: A Table element retrieved from a Table in a Word document containing data.
    :return: A list containing all the data from the table.
    """
    text_input = []

    for row in table.rows: 
        for cell in row.cells:
            paragraphs = []
            for paragraph in cell.paragraphs:
                paragraphs.append(paragraph.text)
            
            text_input.append("\n".join(paragraphs))    # separate each paragraph by a line break

    return text_input

def __checkbox_in_row(row):
    """
    Given a list of data from a row, return TRUE if there is a checkbox in the row, FALSE otherwise.

    :param row: A list of data from a row within a Word document table.
    :return: TRUE if a checkbox is in the row, FALSE otherwise.
    """
    return any([checkbox in row_element for row_element in row for checkbox in ["☒", "☐"]])

def __get_checkbox_index(row, is_checked):
    """
    Given a list of data from a row, return the index of first occurrence the passed check type. Returns -1 if the check type is not present in the list.

    :param row: A list of data from a row within a Word document table.
    :param is_checked: Boolean corresponding to whether the desired index is of checked checkboxes (true) or unchecked checkboxes (unchecked).
    :return: The index of the list that contains the specified checkbox type if it is present. Otherwise, return -1.
    """
    if is_checked:
        checkbox = "☒"
    else:
        checkbox = "☐"

    for id, cell_text in enumerate(row):
        if checkbox in cell_text:
            return id   # returns the ID if the specified checkbox is found
    
    return -1   # returns -1 if the row is looped through and no checkbox is found
